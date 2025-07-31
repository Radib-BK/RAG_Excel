"""
Document ingestion and processing module
Handles text extraction from various file formats and text chunking
"""

import os
import logging
import sqlite3
import uuid
import hashlib
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd

# Document processing imports
import pdfplumber
from docx import Document

# AI/ML imports for better tokenization
try:
    from transformers import AutoTokenizer
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

from utils import (
    extract_text_from_image_file, 
    clean_text, 
    get_file_extension,
    estimate_tokens
)

logger = logging.getLogger(__name__)

class RecursiveTextSplitter:
    """
    LangChain-style recursive text splitter that preserves semantic structure
    Splits on: paragraphs → sentences → words
    """
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Separators in order of preference (most semantic to least)
        self.separators = [
            "\n\n",      # Paragraphs
            "\n",        # Lines  
            ". ",        # Sentences
            "! ",        # Exclamations
            "? ",        # Questions
            "; ",        # Semicolons
            ", ",        # Commas
            " ",         # Words
            ""           # Characters (last resort)
        ]
    
    def split_text(self, text: str) -> List[str]:
        """Recursively split text maintaining semantic boundaries"""
        return self._split_text_recursive(text, self.separators)
    
    def _split_text_recursive(self, text: str, separators: List[str]) -> List[str]:
        """Recursive splitting with fallback to next separator"""
        final_chunks = []
        
        # Base case: no more separators
        if not separators:
            return self._split_by_character(text)
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        # Split by current separator
        splits = text.split(separator) if separator else [text]
        
        current_chunk = ""
        
        for split in splits:
            # If adding this split would exceed chunk size
            if len(current_chunk) + len(split) + len(separator) > self.chunk_size:
                # Save current chunk if it has content
                if current_chunk:
                    final_chunks.append(current_chunk.strip())
                
                # If this split is too large, recursively split it
                if len(split) > self.chunk_size:
                    final_chunks.extend(
                        self._split_text_recursive(split, remaining_separators)
                    )
                    current_chunk = ""
                else:
                    current_chunk = split
            else:
                # Add to current chunk
                if current_chunk:
                    current_chunk += separator + split
                else:
                    current_chunk = split
        
        # Don't forget the last chunk
        if current_chunk:
            final_chunks.append(current_chunk.strip())
        
        # Apply overlap
        return self._add_overlap(final_chunks)
    
    def _split_by_character(self, text: str) -> List[str]:
        """Split by character when all else fails"""
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk = text[i:i + self.chunk_size]
            if chunk.strip():
                chunks.append(chunk)
        return chunks
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between chunks for context continuity"""
        if len(chunks) <= 1:
            return chunks
        
        overlapped_chunks = [chunks[0]]
        
        for i in range(1, len(chunks)):
            # Get overlap from previous chunk
            prev_chunk = chunks[i-1]
            current_chunk = chunks[i]
            
            # Take last N characters from previous chunk as overlap
            if len(prev_chunk) > self.chunk_overlap:
                overlap = prev_chunk[-self.chunk_overlap:]
                # Find a good breaking point (word boundary)
                space_idx = overlap.find(' ')
                if space_idx > 0:
                    overlap = overlap[space_idx:].strip()
                
                overlapped_chunk = overlap + " " + current_chunk
            else:
                overlapped_chunk = current_chunk
            
            overlapped_chunks.append(overlapped_chunk)
        
        return overlapped_chunks

class DocumentProcessor:
    """Handles document processing and text extraction with enhanced capabilities"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.csv', '.db', '.jpg', '.jpeg', '.png']
        self.chunk_size = int(os.getenv('CHUNK_SIZE', 512))
        self.chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 100))
        
        # Initialize recursive text splitter
        self.text_splitter = RecursiveTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
        
        # Initialize real tokenizer if available
        self.tokenizer = None
        if TRANSFORMERS_AVAILABLE:
            try:
                embedding_model = os.getenv('EMBEDDING_MODEL', 'sentence-transformers/all-MiniLM-L6-v2')
                self.tokenizer = AutoTokenizer.from_pretrained(embedding_model)
                logger.info(f"Initialized tokenizer: {embedding_model}")
            except Exception as e:
                logger.warning(f"Could not load tokenizer: {e}")
        
        # Strategy pattern for extractors
        self.extractors = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt,
            'csv': self._extract_from_csv,
            'db': self._extract_from_db,
            'jpg': self._extract_from_image,
            'jpeg': self._extract_from_image,
            'png': self._extract_from_image
        }
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        extension = get_file_extension(filename)
        return f'.{extension}' in self.supported_formats
    
    def extract_text(self, file_path: str, filename: str) -> str:
        """
        Extract text from various file formats using strategy pattern
        
        Args:
            file_path: Path to the file
            filename: Original filename
            
        Returns:
            Extracted text content with enhanced metadata
        """
        extension = get_file_extension(filename)
        
        if extension not in self.extractors:
            raise ValueError(f"Unsupported file format: {extension}")
        
        try:
            return self.extractors[extension](file_path)
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text and tables from PDF using pdfplumber with enhanced structure"""
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_content = []
                    page_content.append(f"[PAGE {page_num + 1}]")
                    
                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        page_content.append(f"TEXT:\n{page_text}")
                    
                    # Extract tables with structure
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            page_content.append(f"\nTABLE {table_idx + 1}:")
                            # Convert table to readable format
                            for row_idx, row in enumerate(table):
                                if row_idx == 0:  # Header row
                                    page_content.append("HEADERS: " + " | ".join([str(cell) if cell else "" for cell in row]))
                                else:
                                    page_content.append("ROW: " + " | ".join([str(cell) if cell else "" for cell in row]))
                    
                    if page_content:
                        text_content.append("\n".join(page_content))
                        
                except Exception as e:
                    logger.warning(f"Error extracting from PDF page {page_num + 1}: {str(e)}")
                    continue
        
        return "\n\n".join(text_content)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text format
            text_content = []
            text_content.append(f"CSV Data with {len(df)} rows and {len(df.columns)} columns")
            text_content.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_content.append("")
            
            # Add sample of data (first 100 rows to avoid huge chunks)
            sample_size = min(100, len(df))
            for idx, row in df.head(sample_size).iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                text_content.append(" | ".join(row_text))
            
            if len(df) > 100:
                text_content.append(f"... and {len(df) - 100} more rows")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading CSV file: {str(e)}")
            raise
    
    def _extract_from_db(self, file_path: str) -> str:
        """Extract text from SQLite database"""
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            
            # Get all table names
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            
            text_content = []
            text_content.append(f"Database contains {len(tables)} tables")
            
            for table_name in tables:
                table_name = table_name[0]
                text_content.append(f"\nTable: {table_name}")
                
                # Get table schema
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = cursor.fetchall()
                column_names = [col[1] for col in columns]
                text_content.append(f"Columns: {', '.join(column_names)}")
                
                # Get sample data (first 50 rows)
                cursor.execute(f"SELECT * FROM {table_name} LIMIT 50")
                rows = cursor.fetchall()
                
                for row in rows:
                    row_text = []
                    for i, value in enumerate(row):
                        if value is not None:
                            row_text.append(f"{column_names[i]}: {value}")
                    text_content.append(" | ".join(row_text))
                
                # Get total count
                cursor.execute(f"SELECT COUNT(*) FROM {table_name}")
                total_rows = cursor.fetchone()[0]
                if total_rows > 50:
                    text_content.append(f"... and {total_rows - 50} more rows in {table_name}")
            
            conn.close()
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading database file: {str(e)}")
            raise
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        return extract_text_from_image_file(file_path)
    
    def _generate_doc_id(self, content: str, filename: str) -> str:
        """Generate unique document ID based on content hash"""
        content_hash = hashlib.md5((content + filename).encode()).hexdigest()
        return f"doc_{content_hash[:12]}"
    
    def _estimate_tokens_accurate(self, text: str) -> int:
        """Accurate token estimation using real tokenizer or fallback"""
        if self.tokenizer:
            try:
                tokens = self.tokenizer.encode(text, add_special_tokens=False)
                return len(tokens)
            except Exception as e:
                logger.debug(f"Tokenizer failed, using fallback: {e}")
        
        # Fallback to original estimation
        return estimate_tokens(text)
    
    def _extract_page_number(self, text: str) -> Optional[int]:
        """Extract page number from text if available"""
        page_match = re.search(r'\[PAGE?\s+(\d+)\]', text, re.IGNORECASE)
        return int(page_match.group(1)) if page_match else None
    
    def chunk_text(self, text: str, source_file: str, source_type: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks with enhanced metadata and semantic preservation
        
        Args:
            text: Text to chunk
            source_file: Source file name for metadata
            source_type: Type of source (pdf, docx, etc.)
            
        Returns:
            List of chunk dictionaries with enhanced metadata
        """
        cleaned_text = clean_text(text)
        
        if not cleaned_text:
            return []
        
        # Generate document ID
        doc_id = self._generate_doc_id(cleaned_text, source_file)
        
        # Use recursive text splitter for semantic chunking
        text_chunks = self.text_splitter.split_text(cleaned_text)
        
        chunks = []
        for chunk_idx, chunk_text in enumerate(text_chunks):
            # Skip very short chunks
            if len(chunk_text.strip()) < 10:
                continue
            
            # Extract page number if available
            page_num = self._extract_page_number(chunk_text)
            
            # Create enhanced metadata
            chunk = {
                'doc_id': doc_id,
                'text': chunk_text.strip(),
                'source': source_file,
                'source_type': source_type,
                'chunk_index': chunk_idx,
                'page': page_num,
                'word_count': len(chunk_text.split()),
                'char_count': len(chunk_text),
                'estimated_tokens': self._estimate_tokens_accurate(chunk_text),
                'has_table': 'TABLE' in chunk_text.upper(),
                'has_figure': 'FIGURE' in chunk_text.upper() or 'IMAGE' in chunk_text.upper(),
                'created_at': pd.Timestamp.now().isoformat() if 'pd' in globals() else None
            }
            
            chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} semantic chunks from {source_file}")
        return chunks
    
    def process_document(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """
        Complete document processing pipeline with enhanced capabilities
        
        Args:
            file_path: Path to the document file
            filename: Original filename
            
        Returns:
            List of processed chunks with enhanced metadata
        """
        logger.info(f"Processing document: {filename}")
        
        # Get source type
        source_type = get_file_extension(filename)
        
        # Extract text with enhanced capabilities
        extracted_text = self.extract_text(file_path, filename)
        
        if not extracted_text.strip():
            logger.warning(f"No text extracted from {filename}")
            return []
        
        logger.info(f"Extracted {len(extracted_text)} characters from {filename}")
        
        # Create semantic chunks with enhanced metadata
        chunks = self.chunk_text(extracted_text, filename, source_type)
        
        # Log chunk statistics
        if chunks:
            avg_tokens = sum(chunk['estimated_tokens'] for chunk in chunks) / len(chunks)
            logger.info(f"Created {len(chunks)} chunks with average {avg_tokens:.1f} tokens each")
            
            # Log special content detection
            table_chunks = sum(1 for chunk in chunks if chunk['has_table'])
            if table_chunks > 0:
                logger.info(f"Detected {table_chunks} chunks with table content")
        
        return chunks
